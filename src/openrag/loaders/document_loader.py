"""Document loaders for various file formats."""

import logging
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from openrag.core.base import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Document loader for various file formats."""

    @staticmethod
    def load_txt(file_path: Path) -> Document:
        """Load text file.

        Args:
            file_path: Path to text file

        Returns:
            Document
        """
        with open(file_path, encoding="utf-8") as f:
            content = f.read()

        return Document(
            content=content,
            metadata={"source": str(file_path), "type": "txt"},
            id=str(file_path),
        )

    @staticmethod
    def load_pdf(file_path: Path) -> Document:
        """Load PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Document
        """
        reader = PdfReader(file_path)
        content_parts = []

        for _page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                content_parts.append(text)

        content = "\n\n".join(content_parts)

        return Document(
            content=content,
            metadata={
                "source": str(file_path),
                "type": "pdf",
                "pages": len(reader.pages),
            },
            id=str(file_path),
        )

    @staticmethod
    def load_docx(file_path: Path) -> Document:
        """Load DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Document
        """
        doc = DocxDocument(file_path)
        content_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        content = "\n\n".join(content_parts)

        return Document(
            content=content,
            metadata={
                "source": str(file_path),
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
            },
            id=str(file_path),
        )

    @staticmethod
    def load_md(file_path: Path) -> Document:
        """Load Markdown file.

        Args:
            file_path: Path to Markdown file

        Returns:
            Document
        """
        with open(file_path, encoding="utf-8") as f:
            content = f.read()

        return Document(
            content=content,
            metadata={"source": str(file_path), "type": "md"},
            id=str(file_path),
        )

    @classmethod
    def load(cls, file_path: Path) -> Document:
        """Load document based on file extension.

        Args:
            file_path: Path to document

        Returns:
            Document

        Raises:
            ValueError: If file format is not supported
        """
        suffix = file_path.suffix.lower()

        loaders = {
            ".txt": cls.load_txt,
            ".pdf": cls.load_pdf,
            ".docx": cls.load_docx,
            ".md": cls.load_md,
        }

        loader = loaders.get(suffix)
        if not loader:
            raise ValueError(f"Unsupported file format: {suffix}")

        logger.info(f"Loading document: {file_path}")
        return loader(file_path)

    @classmethod
    def load_directory(cls, directory: Path, recursive: bool = True) -> list[Document]:
        """Load all supported documents from directory.

        Args:
            directory: Directory path
            recursive: Whether to search recursively

        Returns:
            List of documents
        """
        supported_extensions = {".txt", ".pdf", ".docx", ".md"}
        documents = []

        pattern = "**/*" if recursive else "*"
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    doc = cls.load(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents
